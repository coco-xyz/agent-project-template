"""
Word Parser Tool

Tool for parsing text from Word (.docx) files using multiple methods:
- python-docx library (parse_text_from_word_bytes)
- Direct XML parsing (parse_text_from_word_bytes_xml) for better compatibility
- mammoth library (parse_text_from_word_bytes_mammoth) for most reliable extraction
"""

import io
import zipfile
import xml.etree.ElementTree as ET
from docx import Document
from ai_agents.core.exceptions import DataProcessException
from ai_agents.core.error_codes import DataProcessErrorCode

try:
    import mammoth
except ImportError:
    mammoth = None


def parse_text_from_word_bytes(word_bytes: bytes) -> str:
    """
    Parse text content from Word (.docx) bytes using python-docx.

    Args:
        word_bytes (bytes): Word document content as bytes.

    Returns:
        str: Extracted plain text content from paragraphs and tables.
    """
    try:
        doc = Document(io.BytesIO(word_bytes))
        parts: list[str] = []

        # Extract paragraphs
        parts.extend([p.text for p in doc.paragraphs if p.text])

        # Extract text from tables (if any)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        parts.append(cell.text)

        return "\n".join(parts) if parts else "No text could be extracted from the Word document."
    except Exception as e:
        raise DataProcessException(
            DataProcessErrorCode.DATA_CONVERSION_ERROR,
            detail=f"Error parsing Word document: {str(e)}",
        )


def parse_text_from_word_file(file_path: str) -> str:
    """
    Parse text content from a Word (.docx) file path using python-docx.

    Args:
        file_path (str): Path to the Word document.

    Returns:
        str: Extracted plain text content from paragraphs and tables.
    """
    try:
        doc = Document(file_path)
        parts: list[str] = []

        parts.extend([p.text for p in doc.paragraphs if p.text])

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        parts.append(cell.text)

        return "\n".join(parts) if parts else "No text could be extracted from the Word document."
    except Exception as e:
        raise DataProcessException(
            DataProcessErrorCode.DATA_CONVERSION_ERROR,
            detail=f"Error parsing Word file: {str(e)}",
        )


def parse_text_from_word_bytes_xml(word_bytes: bytes) -> str:
    """
    Parse text content from Word (.docx) bytes using direct XML parsing.

    This method uses simplified XPath queries and adjacent deduplication
    to handle complex DOCX files more effectively.

    Args:
        word_bytes (bytes): Word document content as bytes.

    Returns:
        str: Extracted plain text content from all paragraphs.
    """
    try:
        NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

        with zipfile.ZipFile(io.BytesIO(word_bytes)) as z:
            with z.open("word/document.xml") as f:
                tree = ET.parse(f)

        root = tree.getroot()
        lines = []

        # Extract all paragraphs using simplified XPath
        for p in root.findall(".//w:p", NS):
            texts = [t.text for t in p.findall(".//w:t", NS) if t.text]
            line = "".join(texts).strip()
            if line:
                lines.append(line)

        # 只做相邻去重
        cleaned = []
        for line in lines:
            if not cleaned or cleaned[-1] != line:
                cleaned.append(line)

        return "\n".join(cleaned) if cleaned else "No text could be extracted from the Word document."

    except Exception as e:
        raise DataProcessException(
            DataProcessErrorCode.DATA_CONVERSION_ERROR,
            detail=f"Error parsing Word document with XML method: {str(e)}",
        )


def parse_text_from_word_bytes_mammoth(word_bytes: bytes) -> str:
    """
    Parse text content from Word (.docx) bytes using mammoth library.

    This method provides the most reliable text extraction for DOCX files
    by using mammoth's specialized raw text extraction.

    Args:
        word_bytes (bytes): Word document content as bytes.

    Returns:
        str: Extracted plain text content.
    """
    if mammoth is None:
        raise DataProcessException(
            DataProcessErrorCode.DATA_CONVERSION_ERROR,
            detail="mammoth library not installed. Run: pip install mammoth",
        )

    try:
        result = mammoth.extract_raw_text(io.BytesIO(word_bytes))
        lines = [line.strip() for line in result.value.splitlines() if line.strip()]
        return "\n".join(lines) if lines else "No text could be extracted from the Word document."

    except Exception as e:
        raise DataProcessException(
            DataProcessErrorCode.DATA_CONVERSION_ERROR,
            detail=f"Error parsing Word document with mammoth: {str(e)}",
        )

